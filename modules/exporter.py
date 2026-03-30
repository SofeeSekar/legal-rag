from fpdf import FPDF
from docx import Document
from datetime import datetime
import os, re

EXPORT_DIR = os.path.join(os.path.dirname(__file__), "..", "exports")
os.makedirs(EXPORT_DIR, exist_ok=True)

def _clean(text: str) -> str:
    """Remove markdown symbols for plain text output."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text.strip()

# ── PDF ──────────────────────────────────────────────────────────────────────
class LexPDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(60, 60, 180)
        self.cell(0, 10, "LexRAG — Legal Document Report", align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 8)
        self.set_text_color(140, 140, 140)
        self.cell(0, 6, f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}", align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(4)

    def footer(self):
        self.set_y(-14)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(160, 160, 160)
        self.cell(0, 10, f"Page {self.page_no()} — Confidential", align="C")

    def section(self, title: str):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(60, 60, 180)
        self.ln(4)
        self.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
        self.set_draw_color(60, 60, 180)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(3)

    def body(self, text: str):
        self.set_font("Helvetica", "", 10)
        self.set_text_color(40, 40, 40)
        self.multi_cell(0, 6, _clean(text))
        self.ln(2)

def export_pdf(title: str, sections: list[dict]) -> bytes:
    """sections = [{'title': str, 'content': str}]"""
    pdf = LexPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(20, 20, 60)
    pdf.cell(0, 10, _clean(title), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    for s in sections:
        pdf.section(s["title"])
        pdf.body(s["content"])
    path = os.path.join(EXPORT_DIR, "lexrag_report.pdf")
    pdf.output(path)
    with open(path, "rb") as f:
        return f.read()

# ── WORD ─────────────────────────────────────────────────────────────────────
def export_word(title: str, sections: list[dict]) -> bytes:
    doc = Document()
    doc.add_heading("LexRAG — Legal Document Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}")
    doc.add_heading(_clean(title), level=1)
    for s in sections:
        doc.add_heading(s["title"], level=2)
        doc.add_paragraph(_clean(s["content"]))
    path = os.path.join(EXPORT_DIR, "lexrag_report.docx")
    doc.save(path)
    with open(path, "rb") as f:
        return f.read()
